import os
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Nombre del archivo Word
file_name = "codigo_python.docx"

# 🔹 Si el archivo ya existe, lo eliminamos
if os.path.exists(file_name):
    os.remove(file_name)
    print(f"🗑 Archivo '{file_name}' eliminado.")

# Código en Python que queremos escribir en Word
codigo_python = """# Creamos un directorio sino existe
os.makedirs("/content/drive/My Drive/TFM/datasets/my_dataset/annotations", exist_ok=True)

# Cargamos el dataset json
with open("/content/drive/My Drive/TFM/annotations_coco.json", "r") as f:
    dataset_dicts = json.load(f)

# Dividimos los datos en train y val (e.j., 80% train, 20% val)
train_dicts, val_dicts = train_test_split(dataset_dicts["images"], test_size=0.2, random_state=42)

# Creamos una funcion para redimensionar las imagenes en caso de que sean necesario
def update_image_dimensions(data_list, target_width=1280, target_height=960):
    # Cargamos las imagenes
    for image_entry in data_list:
        image_path = os.path.join("/content/drive/My Drive/TFM/images", image_entry["file_name"])
        img = cv2.imread(image_path)
        if img is not None:
            # Chequeamos si las imagenes necesitan ser redimensionadas
            if img.shape[1] != target_width or img.shape[0] != target_height:
                resized_img = cv2.resize(img, (target_width, target_height))  
                cv2.imwrite(image_path, resized_img)  

            # Actualizamos el JSON con las dimensiones correctas (si han sido redimensioonadas)
            image_entry["height"] = target_height  
            image_entry["width"] = target_width  
        else:
            print(f"Image not found or could not be loaded: {image_path}")


# Actualizamos las imagenes de train y test
update_image_dimensions(train_dicts) 
update_image_dimensions(val_dicts)  

# Create train.json
train_data = {
    "images": train_dicts,
    "annotations": [ann for ann in dataset_dicts["annotations"] if ann["image_id"] in [img["id"] for img in train_dicts]],
    "categories": dataset_dicts["categories"]
}
with open("/content/drive/My Drive/TFM/datasets/my_dataset/annotations/train.json", "w") as f:
    json.dump(train_data, f)

# Creamos Val_data
val_data = {
    "images": val_dicts,
    "annotations": [ann for ann in dataset_dicts["annotations"] if ann["image_id"] in [img["id"] for img in val_dicts]],
    "categories": dataset_dicts["categories"]
}
with open("/content/drive/My Drive/TFM/datasets/my_dataset/annotations/val.json", "w") as f:
    json.dump(val_data, f)

# Esto evita el AssertionError si la celda se ejecuta varias veces
if "suciedad_hands_train" in MetadataCatalog:
    MetadataCatalog.remove("suciedad_hands_train")
if "suciedad_hands_val" in MetadataCatalog:
    MetadataCatalog.remove("suciedad_hands_val")

# Registramos los conjuntos de datos con la imagen_raíz correcta
register_coco_instances("suciedad_hands_train", {}, "/content/drive/My Drive/TFM/datasets/my_dataset/annotations/train.json", "/content/drive/My Drive/TFM/images")
register_coco_instances("suciedad_hands_val", {}, "/content/drive/My Drive/TFM/datasets/my_dataset/annotations/val.json", "/content/drive/My Drive/TFM/images")

# Configuracion y entrenamiento

cfg = get_cfg()
cfg.merge_from_file(model_zoo.get_config_file("COCO-Detection/faster_rcnn_R_50_FPN_3x.yaml"))
cfg.MODEL.WEIGHTS = model_zoo.get_checkpoint_url("COCO-Detection/faster_rcnn_R_50_FPN_3x.yaml")
cfg.DATASETS.TRAIN = ("suciedad_hands_train",)
cfg.DATASETS.TEST = ("suciedad_hands_val",)
cfg.DATALOADER.NUM_WORKERS = 2
cfg.MODEL.WEIGHTS = model_zoo.get_checkpoint_url("COCO-InstanceSegmentation/mask_rcnn_R_50_FPN_3x.yaml")
cfg.SOLVER.IMS_PER_BATCH = 2
cfg.SOLVER.BASE_LR = 0.00025
cfg.SOLVER.MAX_ITER = 1000
cfg.MODEL.ROI_HEADS.BATCH_SIZE_PER_IMAGE = 128
cfg.MODEL.ROI_HEADS.NUM_CLASSES = 2  # suciedad y manos

# Guardamos el output en una carpeta
os.makedirs(cfg.OUTPUT_DIR, exist_ok=True)

trainer = DefaultTrainer(cfg)
trainer.train()
"""

# Crear un nuevo documento de Word
doc = Document()
doc.add_heading('Código en Python', level=1)  # Agregar título

# Agregar el código con formato de texto monoespaciado
code_paragraph = doc.add_paragraph()
run = code_paragraph.add_run(codigo_python)
run.font.name = 'Courier New'  # Fuente monoespaciada
run.font.size = Pt(10)  # Tamaño de fuente

# Aplicar fondo gris al código
shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
code_paragraph._element.get_or_add_pPr().append(shading_elm)

# Guardar el documento
doc.save(file_name)
print(f"✅ Documento Word creado: '{file_name}'")
